#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
政府公文格式整理工具
依据：《党政机关公文格式》GB/T 9704-2012
功能：将 Word 文档按标准公文格式排版
"""

from docx import Document
from docx.shared import Pt, Mm, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys


class GongwenFormatter:
    """公文格式整理器"""

    # 标准页边距（单位：毫米）
    MARGIN_TOP = 37
    MARGIN_BOTTOM = 35
    MARGIN_LEFT = 28
    MARGIN_RIGHT = 26

    # 字体设置
    FONT_TITLE = '方正小标宋简体'
    FONT_BODY = '仿宋_GB2312'
    FONT_HEITI = '黑体'
    FONT_KAITI = '楷体_GB2312'
    FONT_PAGE_NO = '宋体'

    def __init__(self, input_path, output_path=None):
        self.input_path = input_path
        self.output_path = output_path or input_path.replace('.docx', '_规范版.docx')
        self.doc = Document(input_path)

    def setup_page(self):
        """设置页面格式"""
        section = self.doc.sections[0]

        # 设置页边距（转换为厘米）
        section.top_margin = Mm(self.MARGIN_TOP)
        section.bottom_margin = Mm(self.MARGIN_BOTTOM)
        section.left_margin = Mm(self.MARGIN_LEFT)
        section.right_margin = Mm(self.MARGIN_RIGHT)

        # 设置纸张为A4
        section.page_width = Mm(210)
        section.page_height = Mm(297)

        # 设置页眉页脚距离
        section.header_distance = Mm(15)
        section.footer_distance = Mm(15)

    def set_run_font(self, run, font_name, font_size, bold=False):
        """设置字体"""
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def format_title(self, paragraph):
        """格式化标题"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in paragraph.runs:
            self.set_run_font(run, self.FONT_TITLE, 22)  # 2号字 = 22磅

        # 设置段落格式
        paragraph.paragraph_format.line_spacing = Pt(40)  # 固定值40磅
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)

    def format_main_recipient(self, paragraph):
        """格式化主送机关"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for run in paragraph.runs:
            self.set_run_font(run, self.FONT_BODY, 16)  # 3号字 = 16磅

        paragraph.paragraph_format.line_spacing = Pt(28.9)  # 固定值28.9磅
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)  # 顶格

    def format_body_paragraph(self, paragraph, is_first=False):
        """格式化正文段落"""
        text = paragraph.text.strip()

        if not text:
            return

        # 判断标题层级
        level = self.detect_heading_level(text)

        if level == 1:
            # 一级标题：一、
            self.format_heading(paragraph, 1)
        elif level == 2:
            # 二级标题：（一）
            self.format_heading(paragraph, 2)
        elif level == 3:
            # 三级标题：1.
            self.format_heading(paragraph, 3)
        elif level == 4:
            # 四级标题：（1）
            self.format_heading(paragraph, 4)
        else:
            # 普通正文
            self.format_normal_body(paragraph)

    def detect_heading_level(self, text):
        """检测标题层级"""
        # 一级标题：一、二、三、...
        if re.match(r'^[一二三四五六七八九十]+、', text):
            return 1
        # 二级标题：（一）（二）（三）...
        if re.match(r'^（[一二三四五六七八九十]+）', text):
            return 2
        # 三级标题：1. 2. 3.
        if re.match(r'^\d+\.', text):
            return 3
        # 四级标题：（1）（2）（3）...
        if re.match(r'^（\d+）', text):
            return 4
        return 0

    def format_heading(self, paragraph, level):
        """格式化各级标题"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for run in paragraph.runs:
            if level == 1:
                # 一级标题：黑体
                self.set_run_font(run, self.FONT_HEITI, 16, bold=False)
            elif level == 2:
                # 二级标题：楷体
                self.set_run_font(run, self.FONT_KAITI, 16, bold=False)
            elif level == 3:
                # 三级标题：仿宋加粗
                self.set_run_font(run, self.FONT_BODY, 16, bold=True)
            elif level == 4:
                # 四级标题：仿宋
                self.set_run_font(run, self.FONT_BODY, 16, bold=False)

        # 标题段落格式
        paragraph.paragraph_format.line_spacing = Pt(28.9)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)

    def format_normal_body(self, paragraph):
        """格式化普通正文"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in paragraph.runs:
            self.set_run_font(run, self.FONT_BODY, 16)

        # 正文段落格式
        paragraph.paragraph_format.line_spacing = Pt(28.9)  # 固定值28.9磅
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符（约0.74cm）

    def format_attachment_note(self, paragraph):
        """格式化附件说明"""
        text = paragraph.text.strip()
        if '附件' in text and ('：' in text or ':' in text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for run in paragraph.runs:
                self.set_run_font(run, self.FONT_BODY, 16)

            paragraph.paragraph_format.line_spacing = Pt(28.9)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.first_line_indent = Cm(0.74)

    def format_signature(self, paragraphs):
        """格式化落款（署名和日期）"""
        for paragraph in paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            for run in paragraph.runs:
                self.set_run_font(run, self.FONT_BODY, 16)

            paragraph.paragraph_format.line_spacing = Pt(28.9)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.right_indent = Cm(1.48)  # 右空四字（约1.48cm）

    def add_page_numbers(self):
        """添加页码"""
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加页码字段
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        # 设置页码字体
        self.set_run_font(run, self.FONT_PAGE_NO, 14)  # 四号字 = 14磅

    def process(self):
        """处理文档"""
        # 1. 设置页面
        self.setup_page()

        paragraphs = self.doc.paragraphs
        if not paragraphs:
            print("文档为空")
            return

        # 2. 格式化标题（第一段）
        self.format_title(paragraphs[0])

        # 3. 格式化主送机关（第二段，如果有）
        if len(paragraphs) > 1:
            self.format_main_recipient(paragraphs[1])

        # 4. 格式化正文
        for i, para in enumerate(paragraphs[2:], start=2):
            self.format_body_paragraph(para)

        # 5. 格式化附件说明
        for para in paragraphs:
            if '附件' in para.text:
                self.format_attachment_note(para)

        # 6. 格式化落款（最后两段，假设是署名和日期）
        if len(paragraphs) >= 2:
            self.format_signature(paragraphs[-2:])

        # 7. 添加页码
        self.add_page_numbers()

        # 8. 保存
        self.doc.save(self.output_path)
        print(f"已保存到: {self.output_path}")


def main():
    if len(sys.argv) < 2:
        print("用法: python gongwen_formatter.py <输入文件.docx> [输出文件.docx]")
        print("示例: python gongwen_formatter.py 报告.docx")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    formatter = GongwenFormatter(input_file, output_file)
    formatter.process()


if __name__ == '__main__':
    main()
