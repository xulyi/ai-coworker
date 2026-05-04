#!/usr/bin/env python3
"""
Word 三线表生成脚本
==================
根据《学术论文三线表规范指南》生成符合期刊投稿标准的 Word 三线表。

特性：
- 真正的三根黑线：顶线 1.5 磅、中线 0.5 磅、底线 1.5 磅
- 统计符号自动斜体：*t*、*p*、*F*、*U*、*Z*、*H*、*r* 等
- 希腊字母正体：χ²、α、β、η 等
- 表注中的统计符号同样支持斜体
- 支持中文和英文混排字体设置

依赖：
    pip install python-docx

使用示例：
    python word_three_line_table.py

作者：乐义
版本：1.0.0
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


def set_run_font(run, font_name='Times New Roman', font_size=10, bold=False, italic=False):
    """设置字体，支持中英文混排"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'SimSun')  # 中文字体
    rPr.insert(0, rFonts)


def set_cell_border(cell, **kwargs):
    """
    设置单元格边框

    参数：
        top, bottom, left, right = {"sz": 12, "val": "single", "color": "000000", "space": "0"}
        sz: 线宽（1/8磅），12=1.5磅，4=0.5磅
        val: "single"实线, "none"无边框
        color: 颜色代码，黑色="000000"
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge, {"sz": 0, "val": "none", "color": "000000", "space": "0"})
        tag = 'w:{}'.format(edge)
        element = OxmlElement(tag)
        element.set(qn('w:sz'), str(edge_data.get("sz", 0)))
        element.set(qn('w:val'), edge_data.get("val", "none"))
        element.set(qn('w:color'), edge_data.get("color", "000000"))
        element.set(qn('w:space'), str(edge_data.get("space", "0")))
        tcBorders.append(element)

    tcPr.append(tcBorders)


def parse_italic_text(text):
    """
    解析带 *符号* 标记的文本，返回 [(text, italic), ...]

    示例：
        "*t* = 2.34" -> [("t", True), (" = 2.34", False)]
    """
    parts = []
    pattern = r'\*([^*]+)\*'
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            parts.append((text[last_end:match.start()], False))
        parts.append((match.group(1), True))
        last_end = match.end()
    if last_end < len(text):
        parts.append((text[last_end:], False))
    return parts


def add_italic_stat_to_cell(cell, text_parts):
    """
    向单元格添加带斜体统计符号的文本

    参数：
        text_parts: [(text, italic), ...]
        例如：[("t", True), (" = 2.34", False)]
    """
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for text, italic in text_parts:
        run = paragraph.add_run(text)
        set_run_font(run, italic=italic)


def create_word_three_line_table(doc, title, headers, data, note=""):
    """
    创建 Word 三线表（真正的三根黑线）

    参数：
        doc: Document 对象
        title: 表格标题（如"表 1 两组基线特征比较"）
        headers: 表头列表（统计符号需用 *符号* 标记，如 "*t* 值"）
        data: 数据（二维列表）
        note: 表注（支持 *符号* 标记斜体）

    返回：
        table: 生成的表格对象
    """
    # 添加标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    set_run_font(title_run, font_size=10.5, bold=True)

    # 创建表格
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 填充表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        if '*' in header:
            parts = parse_italic_text(header)
            add_italic_stat_to_cell(cell, parts)
        else:
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, bold=True)

    # 填充数据
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if '*' in cell_text:
                parts = parse_italic_text(cell_text)
                add_italic_stat_to_cell(cell, parts)
            else:
                cell.text = str(cell_text)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        set_run_font(run)

    # 设置三线表边框（真正的三根黑线）
    # sz 单位是 1/8 磅：12 = 1.5磅，4 = 0.5磅
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if row_idx == 0:
                # 第一行（表头）：顶线1.5磅 + 底线0.5磅
                set_cell_border(cell,
                    top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                    bottom={"sz": 4, "val": "single", "color": "000000", "space": "0"},
                    left={"sz": 0, "val": "none"},
                    right={"sz": 0, "val": "none"}
                )
            elif row_idx == len(table.rows) - 1:
                # 最后一行：底线1.5磅
                set_cell_border(cell,
                    top={"sz": 0, "val": "none"},
                    bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                    left={"sz": 0, "val": "none"},
                    right={"sz": 0, "val": "none"}
                )
            else:
                # 中间行：无边框
                set_cell_border(cell,
                    top={"sz": 0, "val": "none"},
                    bottom={"sz": 0, "val": "none"},
                    left={"sz": 0, "val": "none"},
                    right={"sz": 0, "val": "none"}
                )

    # 设置列宽
    for col in table.columns:
        col.width = Inches(1.2)

    # 添加表注（表注中的统计符号也支持斜体）
    if note:
        doc.add_paragraph()
        note_para = doc.add_paragraph()
        note_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        parts = parse_italic_text(note)
        for text, italic in parts:
            run = note_para.add_run(text)
            set_run_font(run, font_size=9, italic=italic)

    return table


def create_example_document(output_path='/Users/leyixu/Desktop/三线表示例.docx'):
    """
    创建示例三线表文档，展示三种常见表格类型
    """
    doc = Document()

    # ========== 示例 1：基线特征比较表 ==========
    headers1 = ['变量', '对照组（n=30）', '实验组（n=30）', '统计量', '*p*']
    data1 = [
        ['年龄（岁）', '58.4 ± 9.6', '57.9 ± 8.8', '*t* = 0.21', '.832'],
        ['BMI（kg/m²）', '23.4 ± 3.2', '24.1 ± 2.9', '*t* = 0.90', '.373'],
        ['基线评分', '42.3 ± 8.1', '43.6 ± 7.8', '*t* = 0.64', '.525'],
        ['病程（月）†', '18 [12, 36]', '21 [14, 42]', '*U* = 412', '.485'],
        ['性别（男）', '18 (60.0%)', '16 (53.3%)', 'χ² = 0.27', '.601'],
    ]
    note1 = ("注：连续变量符合正态分布时采用均值 ± 标准差（M ± SD）表示；"
             "分布偏态时采用中位数 [四分位数范围（IQR）] 表示，†为偏态分布。"
             "分类变量采用 n（%）表示。组间差异：连续正态变量采用独立样本 *t* 检验，"
             "偏态变量采用 Mann-Whitney *U* 检验，分类变量采用 χ² 检验。"
             "*p* < .05 视为差异有统计学意义。")

    create_word_three_line_table(
        doc, '表 1 两组受试者基线特征比较（n = 60）',
        headers1, data1, note1
    )

    # 添加空行
    doc.add_paragraph()
    doc.add_paragraph()

    # ========== 示例 2：主要结局比较表（含效应量） ==========
    headers2 = ['指标', '对照组（n=30）', '实验组（n=30）', '统计量', '*p*', "Cohen's d [95% CI]"]
    data2 = [
        ['步行速度（m/s）', '0.82 ± 0.18', '0.97 ± 0.21', '*t* = 2.96', '.004', '0.76 [0.25, 1.27]'],
        ['Berg 平衡量表（分）', '38.4 ± 7.6', '44.2 ± 6.8', '*t* = 3.11', '.003', '0.80 [0.29, 1.31]'],
        ['FIM 运动（分）', '62.3 ± 11.4', '70.8 ± 9.7', '*t* = 3.12', '.003', '0.81 [0.29, 1.32]'],
        ['疼痛 NRS（分）†', '4 [3, 6]', '2 [1, 4]', '*U* = 271', '.012', '*r* = 0.32'],
    ]
    note2 = ("注：连续正态变量采用均值 ± 标准差（M ± SD）表示，组间比较采用独立样本 *t* 检验，"
             "效应量采用 Cohen's d（95% CI）；†疼痛评分不满足正态性假设，采用中位数 [IQR] 表示，"
             "组间比较采用 Mann-Whitney *U* 检验，效应量采用 *r* = |*Z*|/√N。"
             "**p* < .01；所有检验均为双侧检验，α = .05。")

    create_word_three_line_table(
        doc, '表 2 两组受试者干预后主要结局指标比较',
        headers2, data2, note2
    )

    # 添加空行
    doc.add_paragraph()
    doc.add_paragraph()

    # ========== 示例 3：回归分析结果表 ==========
    headers3 = ['变量', '*B*', 'SE', '*β*', '*t*', '*p*', '95% CI']
    data3 = [
        ['常数项', '0.23', '0.18', '—', '1.27', '.208', '[-0.14, 0.60]'],
        ['年龄（岁）', '-0.008', '0.003', '-.284', '-2.66', '.010', '[-0.014, -0.002]'],
        ['BMI（kg/m²）', '-0.012', '0.007', '-.189', '-1.77', '.082', '[-0.026, 0.002]'],
        ['Berg 评分（分）', '0.018', '0.004', '.488', '4.56', '< .001', '[0.010, 0.026]'],
        ['干预组别（实验=1）', '0.142', '0.048', '.315', '2.96', '.004', '[0.046, 0.238]'],
    ]
    note3 = ("注：*B* = 非标准化回归系数；SE = 标准误（正体）；*β* = 标准化回归系数；"
             "95% CI 为 *B* 的置信区间。模型整体：*R²* = .524，调整后 *R²* = .493，"
             "*F* = 15.12，*p* < .001。共线性诊断：所有预测变量 VIF < 3.0，"
             "未发现严重共线性问题。")

    create_word_three_line_table(
        doc, '表 3 以步行速度为结局变量的多元线性回归分析（n = 60）',
        headers3, data3, note3
    )

    # 保存文档
    doc.save(output_path)
    print(f"文档已保存到：{output_path}")
    return output_path


if __name__ == '__main__':
    # 运行示例
    create_example_document()
